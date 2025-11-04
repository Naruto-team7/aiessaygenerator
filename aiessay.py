import streamlit as st
from openai import OpenAI
from PyPDF2 import PdfReader
from docx import Document
from io import BytesIO
import os
import json
from datetime import datetime

# ----------------- PAGE CONFIG -----------------
st.set_page_config(page_title="AI Writing Website", layout="wide")

# ----------------- FILE PATHS -----------------
USER_DATA_FILE = "users.json"
ESSAY_HISTORY_DIR = "essay_history"
os.makedirs(ESSAY_HISTORY_DIR, exist_ok=True)

# ----------------- USER AUTH SYSTEM -----------------
def load_users():
    if not os.path.exists(USER_DATA_FILE):
        return {}
    with open(USER_DATA_FILE, "r") as f:
        return json.load(f)

def save_users(users):
    with open(USER_DATA_FILE, "w") as f:
        json.dump(users, f, indent=4)

users = load_users()

# ----------------- LOGIN / SIGNUP -----------------
st.sidebar.title("🔐 User Login / Signup")
login_mode = st.sidebar.radio("Select Mode", ["Login", "Signup"])
username = st.sidebar.text_input("Username:")
password = st.sidebar.text_input("Password:", type="password")

if login_mode == "Signup":
    if st.sidebar.button("Create Account"):
        if username in users:
            st.sidebar.error("Username already exists.")
        else:
            users[username] = {"password": password}
            save_users(users)
            st.sidebar.success("Account created! Please log in now.")

if login_mode == "Login":
    if st.sidebar.button("Login"):
        if username not in users or users[username]["password"] != password:
            st.sidebar.error("Invalid username or password.")
        else:
            st.session_state["user"] = username
            st.sidebar.success(f"Welcome, {username}!")

# ----------------- AUTH CHECK -----------------
if "user" not in st.session_state:
    st.warning("Please log in to use the AI Writing features.")
    st.stop()

user = st.session_state["user"]
user_history_path = os.path.join(ESSAY_HISTORY_DIR, f"{user}_history.json")

# ----------------- LOAD & SAVE ESSAYS -----------------
def load_history():
    if os.path.exists(user_history_path):
        with open(user_history_path, "r") as f:
            return json.load(f)
    return []

def save_essay(title, content):
    history = load_history()
    history.append({
        "title": title,
        "content": content,
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    })
    with open(user_history_path, "w") as f:
        json.dump(history, f, indent=4)

# ----------------- MAIN INTERFACE -----------------
st.title("📝 AI Writing Website")
st.write(f"Welcome **{user}**! Upload a file or enter a topic to generate your essay.")

# ----------------- OPENAI API SETUP -----------------
api_key = st.secrets.get("OPENAI_API_KEY")
if not api_key:
    st.error("⚠️ Missing API Key! Please add your key in .streamlit/secrets.toml")
    st.stop()

client = OpenAI(api_key=api_key)

# ----------------- FILE UPLOAD -----------------
uploaded_file = st.file_uploader("📄 Upload a PDF or DOCX file", type=["pdf", "docx"])

def extract_text(file):
    text = ""
    if file.name.endswith(".pdf"):
        reader = PdfReader(file)
        for page in reader.pages:
            text += page.extract_text() or ""
    elif file.name.endswith(".docx"):
        doc = Document(file)
        for para in doc.paragraphs:
            text += para.text + "\n"
    return text.strip()

text_data = ""
if uploaded_file is not None:
    text_data = extract_text(uploaded_file)
    st.success("✅ Text extracted successfully!")
    with st.expander("📜 View Extracted Text"):
        st.text_area("Extracted Text", text_data, height=200)

# ----------------- ESSAY SETTINGS -----------------
st.subheader("⚙️ Essay Settings")
topic = st.text_input("Enter your essay topic or prompt:", "")
word_count = st.slider("Select essay length (words)", 500, 3000, 1000, 100)
tone = st.selectbox("Choose essay tone", ["Academic", "Analytical", "Reflective", "Creative", "Persuasive", "Narrative"])

# ----------------- GENERATE ESSAY -----------------
if st.button("🚀 Generate Essay"):
    if not topic and not text_data:
        st.warning("Please provide a topic or upload content first.")
    else:
        with st.spinner("Generating your essay... Please wait ⏳"):
            base_prompt = (
                f"Write a {tone.lower()} essay of around {word_count} words on the topic: '{topic}'. "
                f"Include an introduction, body, and conclusion. Make it coherent, structured, and engaging."
            )
            if text_data:
                base_prompt += "\n\nUse the following text as supporting material:\n" + text_data[:4000]

            try:
                response = client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[{"role": "user", "content": base_prompt}],
                    temperature=0.7
                )
                essay = response.choices[0].message.content.strip()
                st.success("✅ Essay generated successfully!")
                st.text_area("Generated Essay", essay, height=400)

                # Save essay
                save_essay(topic, essay)

                # ----------------- DOWNLOAD OPTIONS -----------------
                st.subheader("⬇️ Download Options")

                # DOCX
                docx_buffer = BytesIO()
                doc = Document()
                doc.add_heading("AI-Generated Essay", 0)
                doc.add_paragraph(essay)
                doc.save(docx_buffer)
                st.download_button(
                    label="📘 Download as DOCX",
                    data=docx_buffer.getvalue(),
                    file_name=f"{topic[:30].replace(' ', '_')}_essay.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

                # TXT
                st.download_button(
                    label="📄 Download as TXT",
                    data=essay.encode("utf-8"),
                    file_name=f"{topic[:30].replace(' ', '_')}_essay.txt",
                    mime="text/plain"
                )

            except Exception as e:
                st.error(f"❌ Error: {e}")

# ----------------- HISTORY SECTION -----------------
st.sidebar.markdown("---")
st.sidebar.header("🕓 Essay History")

if st.sidebar.button("View My Essays"):
    history = load_history()
    if not history:
        st.sidebar.info("No essays found yet.")
    else:
        for entry in reversed(history):
            with st.sidebar.expander(f"{entry['title']} ({entry['timestamp']})"):
                st.sidebar.write(entry["content"][:500] + ("..." if len(entry["content"]) > 500 else ""))
